#!/usr/bin/env python
# -*- coding: utf-8 -*-

import redcap, os
from docx import Document
from requests import post
import dlb

#change API key(second field in redcap.project) for new document generation
project = redcap.Project('https://redcap.vanderbilt.edu/api/', 'F44F337CD4C1B009316F0B2C3E4B913F')

fields ='dl_first_name','dl_last_name','dl_dob','dl_age','dl_doa_45c','dl_sex', \
        'dl_wj_lwid_ss_grade','dl_wj_lwid_pile_grade','dl_wj_wa_ss_grade',\
        'dl_wj_wa_pile_grade','dl_wj_pc_ss_grade','dl_wj_pc_pile_grade', \
        'dl_wj_oc_ss_grade','dl_wj_oc_pile_grade','dl_towre_swe_ss_grade',\
        'dl_towre_swe_pile_grade','dl_towre_pde_ss_grade','dl_towre_pde_pile_grade',\
        'dl_towre_twe_ss_grade','dl_towre_twe_pile_grade','dl_towk_exp_ss',\
        'dl_towk_rec_ss','dl_ctopp_el_ss','dl_ctopp_nr_ss','dl_wiscds_ss'


def document_template_access(doc_path, doc_name):
    """
    Opens the document template used for document generation
    :param doc_name: filename of template document
    :param doc_path: to template document
    :return: True if file exists
    """
    file = os.path.join(doc_path,doc_name)
    if os.path.exists(file):
        document = Document(file)
        return document
    else:
        print ("""Template file doesn't exist: Check filepath and filename provided
        Hint:Also check that suffix of filename is provided eg:docx in filename.docx""")

def redcap_field_fetch(fields):
    """
    Exports data from the redcap form associated with the API Key. 
    Only the fields mentioned are pulled.
    :param fields: fields from the redcap form that are
                   required in the document. By default, it is the global variable 
                   where the fields are mentioned. In case of new document, just change the 
                   fields accordingly
    :return:json object with a list of dicts. Each participant info is a dict
    """
    data = project.export_records(fields=fields, format = 'json')
    return data

def replace_demography(document, data):
    """
    Replaces the demographic information in page 1(Participant ID, dob, age, date
    :param document: docx.Document object. Current copy of document in process for new ID
    :param data: dict structure passed from json object. Contains key value pairs
    :return document: docx.Document object
    """
    for para in document.paragraphs:
        para.text = para.text.replace('[dl_age]',data.get('dl_age')) if not data.get('dl_age')==None \
                else para.text.replace('[dl_age]','NA') #age
        para.text = para.text.replace('[dl_dob]',data.get('dl_dob')) if not data.get('dl_dob')==None \
                else para.text.replace('[dl_dob]','NA') #DOB
        para.text = para.text.replace('[dl_date]',data.get('dl_doa_45c')) if not data.get('dl_doa_45c')==None \
                else para.text.replace('[dl_date]','NA') #date
        para.text = para.text.replace('[dl_first_name]',data.get('dl_first_name')) if not data.get('dl_first_name')==None \
    else para.text.replace('[dl_first_name] [dl_last_name]','ID: '+data.get('dl_participant_id')) #firstname
        para.text = para.text.replace('[dl_last_name]',data.get('dl_last_name')) if not data.get('dl_last_name')==None \
                else para.text.replace('[dl_first_name] [dl_last_name]',data.get('dl_participant_id')) #lastname
    return document

def replace_sex(document, data):
    """
    Replaces his/her with appropriate pronoun based on gender
    :param document: docx.Document object
    :param data: dict structure derived from json list. Contains key value pairs
    :return document: docx.Document object
    """
    for para in document.paragraphs:
        if not data.get('dl_sex')==None:
            para.text = para.text.replace('his/her', 'her') if data.get('dl_sex')=='0' \
                    else para.text.replace('his/her', 'his')
    return document


def replace_table_contents(document, data):
    """
    Replaces the text in WJ-IV tables with actual values
    :param document: docx.Document object
    :param data: dict structure derived from json list. Contains key value pairs
    :return document: docx.Document object
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    #WJIV
                    para.text = para.text.replace('dl_wj_lwid_ss_grade',data.get('dl_wj_lwid_ss_grade')) if not data.get('dl_wj_lwid_ss_grade')==None \
                            else para.text.replace('dl_wj_lwid_ss_grade','NA') #lwid_ss_grade
                    para.text = para.text.replace('dl_wj_lwid_pile_grade',data.get('dl_wj_lwid_pile_grade')) if not data.get('dl_wj_lwid_pile_grade')==None \
                            else para.text.replace('dl_wj_lwid_pile_grade','NA') #lwid_pile_grade
                    para.text = para.text.replace('dl_wj_wa_ss_grade',data.get('dl_wj_wa_ss_grade')) if not data.get('dl_wj_wa_ss_grade')==None \
                            else para.text.replace('dl_wj_wa_ss_grade','NA') #wa_ss_grade
                    para.text = para.text.replace('dl_wj_wa_pile_grade',data.get('dl_wj_wa_pile_grade')) if not data.get('dl_wj_wa_pile_grade')==None \
                            else para.text.replace('dl_wj_wa_pile_grade','NA') #wa_pile_grade
                    para.text = para.text.replace('dl_wj_pc_ss_grade',data.get('dl_wj_pc_ss_grade')) if not data.get('dl_wj_pc_ss_grade')==None \
                            else para.text.replace('dl_wj_pc_ss_grade','NA') #pc_ss_grade
                    para.text = para.text.replace('dl_wj_pc_pile_grade',data.get('dl_wj_pc_pile_grade')) if not data.get('dl_wj_pc_pile_grade')==None \
                            else para.text.replace('dl_wj_pc_pile_grade','NA') #pc_pile_grade
                    para.text = para.text.replace('dl_wj_oc_pile_grade',data.get('dl_wj_oc_pile_grade')) if not data.get('dl_wj_oc_pile_grade')==None \
                            else para.text.replace('dl_wj_oc_pile_grade','NA') #oc_pile_grade
                    para.text = para.text.replace('dl_wj_oc_ss_grade',data.get('dl_wj_oc_ss_grade')) if not data.get('dl_wj_oc_ss_grade')==None \
                            else para.text.replace('dl_wj_oc_ss_grade','NA') #oc_ss_grade
                    #TOWRE-2
                    para.text = para.text.replace('dl_towre_swe_ss_grade',data.get('dl_towre_swe_ss_grade')) if not data.get('dl_towre_swe_ss_grade')==None \
                            else para.text.replace('dl_towre_swe_ss_grade','NA') #swe_ss_grade
                    para.text = para.text.replace('dl_towre_swe_pile_grade',data.get('dl_towre_swe_pile_grade')) if not data.get('dl_towre_swe_pile_grade')==None \
                            else para.text.replace('dl_towre_swe_pile_grade','NA') #swe_pile_grade
                    para.text = para.text.replace('dl_towre_pde_ss_grade',data.get('dl_towre_pde_ss_grade')) if not data.get('dl_towre_pde_ss_grade')==None \
                            else para.text.replace('dl_towre_pde_ss_grade','NA') #pde_ss_grade
                    para.text = para.text.replace('dl_towre_pde_pile_grade',data.get('dl_towre_pde_pile_grade')) if not data.get('dl_towre_pde_pile_grade')==None \
                            else para.text.replace('dl_towre_pde_pile_grade','NA') #pde_pile_grade
                    para.text = para.text.replace('dl_towre_twe_ss_grade',data.get('dl_towre_twe_ss_grade')) if not data.get('dl_towre_twe_ss_grade')==None \
                            else para.text.replace('dl_towre_twe_ss_grade','NA') #twe_ss_grade
                    para.text = para.text.replace('dl_towre_twe_pile_grade',data.get('dl_towre_twe_pile_grade')) if not data.get('dl_towre_twe_pile_grade')==None \
                            else para.text.replace('dl_towre_twe_pile_grade','NA') #twe_pile_grade
                    #TOWK
                    para.text = para.text.replace('dl_towk_exp_ss',data.get('dl_towk_exp_ss')) if not data.get('dl_towk_exp_ss')==None \
                            else para.text.replace('dl_towk_exp_ss','NA') #exp_grade
                    para.text = para.text.replace('dl_towk_rec_ss',data.get('dl_towk_rec_ss')) if not data.get('dl_wj_oc_ss_grade')==None \
                            else para.text.replace('dl_towk_rec_ss','NA') #rec_grade
                    #CTOPP-2
                    para.text = para.text.replace('dl_ctopp_el_ss',data.get('dl_ctopp_el_ss')) if not data.get('dl_ctopp_el_ss')==None \
                            else para.text.replace('dl_ctopp_el_ss','NA') #el_ss
                    para.text = para.text.replace('dl_ctopp_nr_ss',data.get('dl_ctopp_nr_ss')) if not data.get('dl_ctopp_nr_ss')==None \
                            else para.text.replace('dl_ctopp_nr_ss','NA') #nr_ss
                    #WISC-V
                    para.text = para.text.replace('dl_wiscds_ss',data.get('dl_wiscds_ss')) if not data.get('dl_wiscds_ss')==None \
                            else para.text.replace('dl_wiscds_ss','NA') #ss

    return document


if __name__ == '__main__':
    docx = document_template_access('/home/local/VANDERBILT/ramadak/Downloads/','DecLearn Testing Report - Identified w fields.docx')
    data = redcap_field_fetch(fields)
    docx = replace_demography(docx, data[5])
    docx = replace_sex(docx, data[5])
    #print (data[5].get('dl_wj_lwid_ss_grade'))
    docx = replace_table_contents(docx,data[5])
    docx.save('test1.docx')
